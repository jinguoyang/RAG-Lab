"""Verify MarkItDown-backed DOCX/XLSX parsing.

Creates DOCX and XLSX fixtures at runtime so the repository does not need
binary test files for the first MarkItDown integration.
"""

from __future__ import annotations

from pathlib import Path
import sys
from tempfile import TemporaryDirectory

ROOT_DIR = Path(__file__).resolve().parents[2]
BACKEND_DIR = ROOT_DIR / "backend"
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from app.services.document_parsing import parse_document  # noqa: E402


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _create_docx_fixture(path: Path) -> None:
    """Create a minimal DOCX with headings and body text."""
    from docx import Document

    doc = Document()
    doc.add_heading("物料管理规范", level=1)
    doc.add_heading("第一章 总则", level=2)
    doc.add_paragraph("本规范适用于公司所有物料的采购、存储和处理流程。")
    doc.add_heading("第二章 物料分类", level=2)
    doc.add_paragraph("物料分为原材料、半成品和成品三大类。")
    doc.save(path)


def _create_xlsx_fixture(path: Path) -> None:
    """Create a minimal workbook with headers and business text."""
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "物料清单"
    sheet.append(["物料编码", "物料名称", "处理策略"])
    sheet.append(["M-001", "呆滞电源模块", "优先调拨"])
    sheet.append(["M-002", "备品电缆", "季度复核"])
    workbook.save(path)


def verify_docx_uses_markdown_chunks() -> None:
    """DOCX should parse through MarkItDown and preserve meaningful text."""
    with TemporaryDirectory() as tmp_dir:
        docx_path = Path(tmp_dir) / "markitdown-sample.docx"
        _create_docx_fixture(docx_path)
        parsed = parse_document(docx_path.name, None, docx_path.read_bytes())

    contents = "\n".join(block.content for block in parsed.blocks)
    _assert(parsed.parser_name == "markitdown_docx", f"DOCX 应使用 markitdown_docx parserName，实际: {parsed.parser_name}")
    _assert(parsed.blocks, "DOCX 未生成 Block")
    _assert(any(block.section for block in parsed.blocks), "DOCX Block 应包含章节或段落定位")
    _assert("物料" in contents or "规范" in contents, "DOCX 样例核心文本未进入 Block")


def verify_xlsx_uses_markdown_chunks() -> None:
    """XLSX should parse table-like content into searchable blocks."""
    with TemporaryDirectory() as tmp_dir:
        xlsx_path = Path(tmp_dir) / "markitdown-sample.xlsx"
        _create_xlsx_fixture(xlsx_path)
        parsed = parse_document(xlsx_path.name, None, xlsx_path.read_bytes())

    contents = "\n".join(block.content for block in parsed.blocks)
    _assert(parsed.parser_name == "markitdown_xlsx", f"XLSX 应使用 markitdown_xlsx parserName，实际: {parsed.parser_name}")
    _assert(parsed.blocks, "XLSX 未生成 Block")
    _assert("物料编码" in contents, "XLSX 表头未进入 Block")
    _assert("呆滞电源模块" in contents, "XLSX 单元格内容未进入 Block")
    _assert(all(block.metadata.get("sourceExtension") == ".xlsx" for block in parsed.blocks), "XLSX metadata 扩展名错误")


def main() -> None:
    verify_docx_uses_markdown_chunks()
    verify_xlsx_uses_markdown_chunks()
    print("MarkItDown DOCX/XLSX parse verification passed.")


if __name__ == "__main__":
    main()
